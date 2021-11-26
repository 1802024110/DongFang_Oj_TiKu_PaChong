from docx import Document
from docx.shared import Inches
import get_page_html
import get_answer
document = Document()
def write_text(title,leve,text,size):
  document.add_heading(title, leve)
  paragraph = document.add_paragraph(text)
  paragraph_format = paragraph.paragraph_format
  paragraph_format.first_line_indent = Inches(size)
def create(start,end,username='',password=''):
  for item in range(start,end+1):
      get_answer.login(username,password)
      document.add_heading(str(item)+'——'+'入门', 1)
      
      write_text('题目',2,get_page_html.get_page_title(item),0.25)
      write_text('题目描述',2,get_page_html.get_page_describe(item),0.25)
      write_text('输入',2,get_page_html.get_page_input_text(item),0.25)
      write_text('输出',2,get_page_html.get_page_output_text(item),0.25)
      write_text('样例输入',2,get_page_html.get_page_Sample_input_text(item),0.25)
      write_text('样例输出',2,get_page_html.get_page_Sample_output_text(item),0.25)
      write_text('提示',2,get_page_html.get_page_Tips(item),0.25)
      write_text('来源',2,get_page_html.get_page_source(item),0.25)
      write_text('解析',2,get_answer.get_answer_text(item),0.25)
      document.add_page_break()#分页符
      print('成功',item)
  document.save('Oj.docx')